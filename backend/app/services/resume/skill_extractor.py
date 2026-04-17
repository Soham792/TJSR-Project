"""
Resume text extraction and skill detection.

Uses:
  - pypdf  for PDF files
  - python-docx for DOCX files
  - keyword matching against a comprehensive tech-skill taxonomy
"""

from __future__ import annotations
import io
import re
import logging
from typing import Sequence

logger = logging.getLogger(__name__)

# ── Comprehensive skills taxonomy ────────────────────────────────────────────
# Each entry is (display_name, [aliases / alternate spellings])
# Matching is case-insensitive and uses word-boundary checks.
_SKILLS_TAXONOMY: list[tuple[str, list[str]]] = [
    # Languages
    ("Python",        ["python"]),
    ("JavaScript",    ["javascript", "js"]),
    ("TypeScript",    ["typescript", "ts"]),
    ("Java",          ["java"]),
    ("C++",           ["c\\+\\+", "cpp"]),
    ("C#",            ["c#", "csharp", "c sharp"]),
    ("Go",            ["golang", r"\bgo\b"]),
    ("Rust",          ["rust"]),
    ("Ruby",          ["ruby"]),
    ("PHP",           ["php"]),
    ("Swift",         ["swift"]),
    ("Kotlin",        ["kotlin"]),
    ("R",             [r"\bR\b"]),
    ("Scala",         ["scala"]),
    ("MATLAB",        ["matlab"]),
    ("Dart",          ["dart"]),
    ("Perl",          ["perl"]),
    ("Shell",         ["bash", "shell scripting", "zsh"]),
    ("SQL",           [r"\bSQL\b"]),
    ("HTML",          ["html"]),
    ("CSS",           [r"\bCSS\b"]),
    ("SASS",          ["sass", "scss"]),
    # Web Frameworks / Libraries
    ("React",         ["react", "react.js", "reactjs"]),
    ("Next.js",       ["next.js", "nextjs"]),
    ("Vue.js",        ["vue", "vue.js", "vuejs"]),
    ("Nuxt.js",       ["nuxt", "nuxt.js"]),
    ("Angular",       ["angular", "angularjs"]),
    ("Svelte",        ["svelte", "sveltekit"]),
    ("jQuery",        ["jquery"]),
    ("Express.js",    ["express", "express.js"]),
    ("Node.js",       ["node", "node.js", "nodejs"]),
    ("NestJS",        ["nestjs", "nest.js"]),
    ("Django",        ["django"]),
    ("Flask",         ["flask"]),
    ("FastAPI",       ["fastapi"]),
    ("Spring",        ["spring boot", "spring framework"]),
    ("Laravel",       ["laravel"]),
    ("Rails",         ["ruby on rails", r"\brails\b"]),
    ("ASP.NET",       ["asp.net", "aspnet", ".net core", "dotnet"]),
    ("GraphQL",       ["graphql"]),
    ("REST API",      ["rest api", "restful", "rest services"]),
    ("gRPC",          ["grpc"]),
    # Databases
    ("PostgreSQL",    ["postgresql", "postgres"]),
    ("MySQL",         ["mysql"]),
    ("MongoDB",       ["mongodb", "mongo"]),
    ("Redis",         ["redis"]),
    ("Elasticsearch", ["elasticsearch", "elastic"]),
    ("Cassandra",     ["cassandra"]),
    ("DynamoDB",      ["dynamodb"]),
    ("SQLite",        ["sqlite"]),
    ("Oracle DB",     ["oracle database", r"\boracle\b"]),
    ("SQL Server",    ["sql server", "mssql", "microsoft sql"]),
    ("Neo4j",         ["neo4j"]),
    ("Supabase",      ["supabase"]),
    ("Firebase",      ["firebase", "firestore"]),
    ("Snowflake",     ["snowflake"]),
    ("BigQuery",      ["bigquery"]),
    # Cloud & Infrastructure
    ("AWS",           ["aws", "amazon web services", "amazon s3", "ec2", "lambda", "cloudwatch"]),
    ("Azure",         ["azure", "microsoft azure"]),
    ("GCP",           ["gcp", "google cloud", "google cloud platform"]),
    ("Docker",        ["docker"]),
    ("Kubernetes",    ["kubernetes", "k8s"]),
    ("Terraform",     ["terraform"]),
    ("Ansible",       ["ansible"]),
    ("Helm",          ["helm"]),
    ("Jenkins",       ["jenkins"]),
    ("GitHub Actions",["github actions"]),
    ("GitLab CI",     ["gitlab ci", "gitlab-ci"]),
    ("CircleCI",      ["circleci"]),
    ("ArgoCD",        ["argocd", "argo cd"]),
    # ML / AI / Data
    ("TensorFlow",    ["tensorflow"]),
    ("PyTorch",       ["pytorch"]),
    ("Scikit-learn",  ["scikit-learn", "sklearn"]),
    ("Keras",         ["keras"]),
    ("Hugging Face",  ["hugging face", "huggingface", "transformers"]),
    ("OpenAI",        ["openai", "gpt-4", "gpt4", "chatgpt"]),
    ("LangChain",     ["langchain"]),
    ("LlamaIndex",    ["llamaindex", "llama index"]),
    ("Pandas",        ["pandas"]),
    ("NumPy",         ["numpy"]),
    ("Matplotlib",    ["matplotlib"]),
    ("Seaborn",       ["seaborn"]),
    ("Apache Spark",  ["spark", "pyspark", "apache spark"]),
    ("Hadoop",        ["hadoop", "hdfs", "mapreduce"]),
    ("Kafka",         ["kafka", "apache kafka"]),
    ("Airflow",       ["airflow", "apache airflow"]),
    ("dbt",           [r"\bdbt\b"]),
    ("Power BI",      ["power bi", "powerbi"]),
    ("Tableau",       ["tableau"]),
    ("Looker",        ["looker"]),
    # Version Control & Collaboration
    ("Git",           [r"\bgit\b"]),
    ("GitHub",        ["github"]),
    ("GitLab",        ["gitlab"]),
    ("Jira",          ["jira"]),
    ("Confluence",    ["confluence"]),
    ("Notion",        ["notion"]),
    # Security & Auth
    ("OAuth",         ["oauth", "oauth2"]),
    ("JWT",           [r"\bjwt\b"]),
    ("SSL/TLS",       ["ssl", "tls"]),
    # Practices & Soft-tech
    ("Agile",         ["agile"]),
    ("Scrum",         ["scrum"]),
    ("Microservices", ["microservices", "micro-services"]),
    ("CI/CD",         ["ci/cd", "continuous integration", "continuous deployment", "continuous delivery"]),
    ("DevOps",        ["devops"]),
    ("Machine Learning",["machine learning", r"\bml\b"]),
    ("Deep Learning", ["deep learning"]),
    ("NLP",           [r"\bnlp\b", "natural language processing"]),
    ("Computer Vision",["computer vision", r"\bcv\b"]),
    ("Data Science",  ["data science", "data scientist"]),
    ("Data Engineering",["data engineering", "data engineer"]),
    ("Cloud Computing",["cloud computing"]),
    ("System Design", ["system design", "distributed systems"]),
    ("Blockchain",    ["blockchain", "solidity", "web3"]),
    # Testing
    ("Jest",          [r"\bjest\b"]),
    ("Pytest",        ["pytest"]),
    ("Selenium",      ["selenium"]),
    ("Cypress",       ["cypress"]),
    # Mobile
    ("React Native",  ["react native"]),
    ("Flutter",       ["flutter"]),
    ("Android",       ["android"]),
    ("iOS",           [r"\bios\b"]),
    # Tools
    ("Linux",         ["linux", "ubuntu", "debian", "centos"]),
    ("Nginx",         ["nginx"]),
    ("RabbitMQ",      ["rabbitmq"]),
    ("Celery",        ["celery"]),
    ("Webpack",       ["webpack"]),
    ("Vite",          ["vite"]),
    ("Tailwind CSS",  ["tailwind", "tailwindcss"]),
]

# Pre-compile patterns for performance
_COMPILED: list[tuple[str, list[re.Pattern]]] = []
for _display, _aliases in _SKILLS_TAXONOMY:
    _patterns = []
    for _alias in _aliases:
        # Wrap in word boundaries unless the alias already contains \b
        if "\\b" in _alias or "\\+" in _alias:
            pat = _alias
        else:
            pat = r"\b" + re.escape(_alias) + r"\b"
        _patterns.append(re.compile(pat, re.IGNORECASE))
    _COMPILED.append((_display, _patterns))


def extract_skills_from_text(text: str) -> list[str]:
    """Return a deduplicated list of skills found in *text*."""
    found: list[str] = []
    for display_name, patterns in _COMPILED:
        for pattern in patterns:
            if pattern.search(text):
                found.append(display_name)
                break  # move to next skill once matched
    return found


# ── File parsing ─────────────────────────────────────────────────────────────

def extract_text_from_pdf(data: bytes) -> str:
    """Extract plain text from PDF bytes using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except ImportError:
        logger.warning("pypdf not installed; cannot parse PDF")
        return ""
    except Exception as exc:
        logger.warning(f"PDF parsing error: {exc}")
        return ""


def extract_text_from_docx(data: bytes) -> str:
    """Extract plain text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [para.text for para in doc.paragraphs]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except ImportError:
        logger.warning("python-docx not installed; cannot parse DOCX")
        return ""
    except Exception as exc:
        logger.warning(f"DOCX parsing error: {exc}")
        return ""


def extract_text_from_txt(data: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_resume(filename: str, data: bytes) -> tuple[str, list[str]]:
    """
    Parse uploaded resume file → (extracted_text, skill_list).
    Supports PDF, DOCX, DOC (as plain text), and TXT.
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = extract_text_from_pdf(data)
    elif lower.endswith((".docx", ".doc")):
        text = extract_text_from_docx(data)
    else:
        text = extract_text_from_txt(data)

    skills = extract_skills_from_text(text)
    return text, skills
